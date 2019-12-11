import os
import docx

techSummary = {}
fileNames = []
for root, dirs, files in os.walk("."):
    for filename in files:
        if filename.endswith("Resume.docx"):
            fileNames.append(filename)

print(fileNames)

for resume in fileNames:
    doc = docx.Document(resume)
    for para_index in range(len(doc.paragraphs)):
        # print(para.text)
        if doc.paragraphs[para_index].text.lower() == 'technology summary':
            var_string = ''
            for string in doc.paragraphs[para_index + 1].runs:
                var_string += string.text
            techSummary[resume] = var_string

print(techSummary)
